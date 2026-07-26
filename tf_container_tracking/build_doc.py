"""Build a professional stakeholder-ready Word document from the scope plan."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

IMG_DIR = os.path.join(os.path.dirname(__file__), "docs", "images")

# ── Colour palette (R, G, B tuples) ──────────────────────────────────────────
DARK_NAVY   = (0x1A, 0x2E, 0x4A)
ACCENT_BLUE = (0x1F, 0x6F, 0xEB)
LIGHT_BLUE  = (0xE8, 0xF1, 0xFD)
MID_GREY    = (0x5A, 0x6A, 0x7A)
WHITE       = (0xFF, 0xFF, 0xFF)
AMBER       = (0xFF, 0xB8, 0x00)
AMBER_BG    = (0xFF, 0xF8, 0xE1)
AMBER_DARK  = (0x7A, 0x50, 0x00)
SKY         = (0xA0, 0xC8, 0xFF)
STEEL       = (0xB0, 0xC8, 0xE0)
SLATE       = (0x70, 0x90, 0xB0)

def c(t):
    """Convert (R,G,B) tuple to RGBColor."""
    return RGBColor(t[0], t[1], t[2])

def h(t):
    """Convert (R,G,B) tuple to hex string."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h(color))
    tcPr.append(shd)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, spec in edges.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in spec.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color=None):
    if color is None:
        color = ACCENT_BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), h(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    p.add_run()._r.append(br)


# ── Styled paragraph helpers ──────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    run.font.color.rgb = c(DARK_NAVY)
    run.font.size = Pt(15)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    run.font.color.rgb = c(ACCENT_BLUE)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def bullet(doc, text, level=0, bold_parts=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    _add_text_with_bold(p, text, bold_parts)
    return p


def numbered_item(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


def embed_image(doc, filename, caption_text=None):
    """Embed an image from IMG_DIR, centred, max width 6 inches."""
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(6.0))
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(caption_text)
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = c(MID_GREY)


def add_section_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(8)
    run.font.color.rgb = c(MID_GREY)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    return p


def _add_text_with_bold(p, text, bold_parts=None):
    if not bold_parts:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        return
    remaining = text
    for frag in bold_parts:
        idx = remaining.find(frag)
        if idx == -1:
            continue
        if idx > 0:
            r = p.add_run(remaining[:idx])
            r.font.size = Pt(10.5)
        br = p.add_run(frag)
        br.bold = True
        br.font.size = Pt(10.5)
        remaining = remaining[idx + len(frag):]
    if remaining:
        r = p.add_run(remaining)
        r.font.size = Pt(10.5)


# ── Decision callout box ──────────────────────────────────────────────────────

def decision_box(doc, items):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, AMBER_BG)
    set_cell_border(cell,
        top={"val": "single", "sz": "12", "color": h(AMBER)},
        bottom={"val": "single", "sz": "6", "color": h(AMBER)},
        left={"val": "single", "sz": "24", "color": h(AMBER)},
        right={"val": "single", "sz": "6", "color": h(AMBER)},
    )
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    lbl_p = cell.add_paragraph()
    lbl_r = lbl_p.add_run("  Required Decisions")
    lbl_r.bold = True
    lbl_r.font.size = Pt(10)
    lbl_r.font.color.rgb = c(AMBER_DARK)
    lbl_p.paragraph_format.space_before = Pt(4)
    lbl_p.paragraph_format.space_after = Pt(2)

    for i, item in enumerate(items, 1):
        dp = cell.add_paragraph(style="List Number")
        dr = dp.add_run(item)
        dr.font.size = Pt(10)
        dr.font.color.rgb = c((0x3A, 0x2A, 0x00))
        dp.paragraph_format.space_before = Pt(1)
        dp.paragraph_format.space_after = Pt(3) if i == len(items) else Pt(1)

    spacer(doc, 4)


# ── Cover page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, DARK_NAVY)

    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    def cp(txt, size, color, bold=False, indent=0.4, sb=0, sa=0):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        r = p.add_run(txt)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = c(color)
        return p

    cp("", 4, WHITE, sb=30)
    cp("Dispatch & Transfer", 28, WHITE, bold=True, sb=0, sa=4)
    cp("Enhancement Plan", 22, SKY, sb=0, sa=0)
    cp("━━━━━━━━━━━", 14, ACCENT_BLUE, sb=10, sa=0)
    cp(f"Scope Proposal  ·  {date.today().strftime('%B %d, %Y')}", 10, STEEL, sb=8, sa=0)
    cp("CONFIDENTIAL — FOR STAKEHOLDER REVIEW", 8, SLATE, bold=True, sb=4, sa=30)

    spacer(doc, 10)

    # Summary inset
    sum_tbl = doc.add_table(rows=1, cols=1)
    sc = sum_tbl.cell(0, 0)
    set_cell_bg(sc, LIGHT_BLUE)
    set_cell_border(sc,
        left={"val": "single", "sz": "18", "color": h(ACCENT_BLUE)},
        top={"val": "none", "sz": "0", "color": "auto"},
        bottom={"val": "none", "sz": "0", "color": "auto"},
        right={"val": "none", "sz": "0", "color": "auto"},
    )
    sc.paragraphs[0]._element.getparent().remove(sc.paragraphs[0]._element)

    intro_p = sc.add_paragraph()
    intro_p.paragraph_format.left_indent = Inches(0.15)
    intro_p.paragraph_format.space_before = Pt(8)
    intro_p.paragraph_format.space_after = Pt(4)
    ir = intro_p.add_run(
        "This document defines the proposed dispatch and transfer enhancements requested by "
        "the client. The objective is to align all stakeholders on scope, sequence, and "
        "delivery expectations before execution."
    )
    ir.font.size = Pt(10.5)
    ir.font.color.rgb = c(DARK_NAVY)

    for line in [
        "Functional requirements and implementation approach",
        "User interface placement and annotated visual references",
        "Decision points, effort estimate, and delivery timeline",
    ]:
        bp = sc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.35)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(1)
        br_ = bp.add_run(line)
        br_.font.size = Pt(10)
        br_.font.color.rgb = c(DARK_NAVY)

    sc.add_paragraph().paragraph_format.space_after = Pt(6)

    spacer(doc, 6)
    page_break(doc)


# ── Scope summary table ───────────────────────────────────────────────────────

def build_scope_summary(doc):
    heading1(doc, "1. Scope Summary")
    add_horizontal_rule(doc)
    body(doc, "The proposed scope includes five functional areas:")

    areas = [
        ("Dispatch Ticket Management Workflow",
         "Dedicated dispatch process to coordinate drivers and equipment with lifecycle state management."),
        ("WhatsApp Dispatch Communication Tracking (Manual)",
         "Controlled step to confirm driver instructions were sent via WhatsApp, with full audit trail."),
        ("Trailer Resource and Location Tracking",
         "Visibility into trailer availability and last known location, updated on dispatch completion."),
        ("Manual Dispatch / Route Creation from Workflow Actions",
         "Replace automatic route generation with explicit user-driven workflow action buttons."),
        ("Sales Order-Based Serial Filtering in Internal Transfers",
         "Constrain serial selection to serials linked to the selected Sales Order."),
    ]

    tbl = doc.add_table(rows=len(areas) + 1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for cell, txt in zip(tbl.rows[0].cells, ["#", "Area", "Summary"]):
        set_cell_bg(cell, ACCENT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.bold = True
        r.font.color.rgb = c(WHITE)
        r.font.size = Pt(10)

    tbl.columns[0].width = Inches(0.4)
    tbl.columns[1].width = Inches(2.4)
    tbl.columns[2].width = Inches(3.7)

    for i, (area, summary) in enumerate(areas, 1):
        row = tbl.rows[i]
        bg = LIGHT_BLUE if i % 2 == 0 else WHITE

        num_c = row.cells[0]
        set_cell_bg(num_c, bg)
        num_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        np_ = num_c.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np_.add_run(str(i))
        nr.font.bold = True
        nr.font.size = Pt(10)
        nr.font.color.rgb = c(ACCENT_BLUE)

        area_c = row.cells[1]
        set_cell_bg(area_c, bg)
        area_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ar = area_c.paragraphs[0].add_run(area)
        ar.font.bold = True
        ar.font.size = Pt(10)

        sum_c = row.cells[2]
        set_cell_bg(sum_c, bg)
        sum_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sr = sum_c.paragraphs[0].add_run(summary)
        sr.font.size = Pt(10)

    spacer(doc, 8)


# ── Generic requirement section ───────────────────────────────────────────────

def req_section(doc, num, title, biz_intro, biz_bullets, impl_items, ui_items, decisions,
               images=None):
    """images: list of (filename, caption) tuples to embed after UI Placement."""
    page_break(doc)
    add_section_label(doc, f"Requirement {num}")
    heading1(doc, f"{num + 2}. Requirement {num}: {title}")
    add_horizontal_rule(doc)

    heading2(doc, f"{num + 2}.1  Business Requirement")
    if biz_intro:
        body(doc, biz_intro)
    for b in biz_bullets:
        bullet(doc, b)

    heading2(doc, f"{num + 2}.2  Proposed Implementation")
    for item in impl_items:
        if isinstance(item, dict):
            bullet(doc, item["text"], bold_parts=item.get("bold"))
            for sub in item.get("subs", []):
                bullet(doc, sub, level=1)
        elif isinstance(item, tuple):
            bullet(doc, item[0], bold_parts=item[1] if len(item) > 1 else None)
        else:
            bullet(doc, item)

    heading2(doc, f"{num + 2}.3  UI Placement")
    for u in ui_items:
        bullet(doc, u)

    if images:
        heading2(doc, f"{num + 2}.4  Visual Reference")
        for fname, caption in images:
            embed_image(doc, fname, caption)

    decision_box(doc, decisions)


# ── Effort breakdown table ────────────────────────────────────────────────────

def build_effort_table(doc):
    rows_data = [
        ("Dispatch Ticket workflow (data model, lifecycle, history)", "18", "6", "2", "26"),
        ("WhatsApp manual tracking (button + sent audit fields)", "4", "2", "1", "7"),
        ("Trailer tracking and location update behaviour", "8", "3", "1", "12"),
        ("Manual workflow actions (replace auto route creation)", "10", "4", "2", "16"),
        ("Internal transfer SO-based serial filtering", "12", "6", "2", "20"),
        ("Deployment, documentation, and handover", "6", "2", "2", "10"),
    ]
    totals = ("Total", "58", "23", "10", "91 hrs")
    headers = ["Workstream", "Build (hrs)", "QA (hrs)", "UAT (hrs)", "Total (hrs)"]

    tbl = doc.add_table(rows=len(rows_data) + 2, cols=5)
    tbl.style = "Table Grid"

    widths = [Inches(3.0), Inches(0.9), Inches(0.8), Inches(0.9), Inches(0.85)]
    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # Header
    for j, (cell, txt) in enumerate(zip(tbl.rows[0].cells, headers)):
        set_cell_bg(cell, ACCENT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(txt)
        r.font.bold = True
        r.font.color.rgb = c(WHITE)
        r.font.size = Pt(9.5)

    # Data rows
    for i, row_data in enumerate(rows_data, 1):
        bg = LIGHT_BLUE if i % 2 == 0 else WHITE
        for j, (cell, txt) in enumerate(zip(tbl.rows[i].cells, row_data)):
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt)
            r.font.size = Pt(9.5)

    # Totals row
    for j, (cell, txt) in enumerate(zip(tbl.rows[-1].cells, totals)):
        set_cell_bg(cell, DARK_NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(txt)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = c(WHITE)

    spacer(doc, 6)


# ── Phased delivery cards ─────────────────────────────────────────────────────

def build_phases(doc):
    heading1(doc, "9. Phased Delivery Plan")
    add_horizontal_rule(doc)

    phases = [
        ("Phase 1", "Core Dispatch Rollout", "6 to 8 working days", DARK_NAVY, WHITE, SKY, AMBER,
         ["Dispatch ticket workflow", "Trailer tracking", "Manual workflow action buttons"]),
        ("Phase 2", "Transfer Control Enhancements", "3 to 4 working days", ACCENT_BLUE, WHITE, SKY, AMBER,
         ["Sales Order-based serial filtering",
          "Validation and edge-case handling for transfer serial availability"]),
        ("Phase 3", "Deployment and UAT Closure", "2 to 3 working days", LIGHT_BLUE, DARK_NAVY,
         ACCENT_BLUE, AMBER_DARK,
         ["Controlled deployment", "Stakeholder walkthrough", "UAT support and closure updates"]),
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    for i, (phase, subtitle, duration, bg, fg, sub_col, dur_col, items) in enumerate(phases):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, bg)
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

        def add(txt, size, color, bold=False, indent=0.1, sb=0, sa=0):
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after = Pt(sa)
            r = p.add_run(txt)
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = c(color)

        add(phase, 11, fg, bold=True, sb=8, sa=2)
        add(subtitle, 9, sub_col, sb=0, sa=4)
        add(f"Est. {duration}", 8.5, dur_col, bold=True, sb=0, sa=6)

        for item in items:
            bp = cell.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            br_ = bp.add_run(item)
            br_.font.size = Pt(9)
            br_.font.color.rgb = c(fg)

        cell.add_paragraph().paragraph_format.space_after = Pt(8)

    spacer(doc, 8)


# ── Main builder ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    build_cover(doc)
    build_scope_summary(doc)

    # ── Requirement 1: Dispatch Ticket Workflow
    req_section(doc, 1, "Dispatch Ticket Workflow",
        biz_intro="The client requires a dedicated dispatch process to coordinate drivers and equipment. "
                  "Each dispatch record must include:",
        biz_bullets=[
            "Sales Order", "Container (optional)", "Location", "Truck",
            "Driver", "Trailer", "WhatsApp Sent", "Dispatch Date", "Completed Status",
        ],
        impl_items=[
            "Introduce a new dispatch ticket entity and associated list / form views.",
            {"text": "Implement controlled lifecycle states:", "subs": [
                "Draft", "Sent", "In Progress", "Completed",
            ]},
            {"text": "On completion:", "subs": [
                "Record completion timestamp and user.",
                "Retain full history in chatter / audit trail.",
                "Prevent unintended post-completion edits (role-based control).",
            ]},
            "Maintain direct linkage to Sales Order and (when applicable) Container.",
        ],
        ui_items=[
            "Inventory main area",
            "New Dispatch menu and sub-screens",
            "Linked access from container workflow records",
        ],
        decisions=[
            "Should one dispatch ticket represent exactly one trip?",
            "After completion, should editing be restricted to managers only?",
        ],
        images=[
            ("03_container_form_proposed_dispatch_buttons.png",
             "Proposed dispatch buttons and summary on container form"),
        ],
    )

    # ── Requirement 2: WhatsApp Communication
    req_section(doc, 2, "WhatsApp Dispatch Communication (Manual Mode)",
        biz_intro="Dispatchers must have a controlled step to confirm that driver instructions "
                  "were sent through WhatsApp.",
        biz_bullets=[],
        impl_items=[
            "The approved approach is manual WhatsApp tracking — no API automation in this scope.",
            ("Add a Send WhatsApp action in the dispatch flow.", ["Send WhatsApp"]),
            {"text": "Capture and store:", "subs": [
                "WhatsApp Sent (Yes / No)",
                "Sent On (timestamp)",
                "Sent By (user)",
            ]},
            "Keep this information visible in dispatch records for operations follow-up.",
        ],
        ui_items=["Dispatch Ticket form view under Inventory > Dispatch"],
        decisions=[
            'Should dispatch move to "Sent" automatically when "Send WhatsApp" is clicked?',
        ],
        images=[
            ("03_container_form_proposed_dispatch_buttons.png",
             "Proposed WhatsApp tracking placement on dispatch form"),
        ],
    )

    # ── Requirement 3: Trailer Resource Tracking
    req_section(doc, 3, "Trailer Resource Tracking",
        biz_intro="The client has limited trailer capacity and needs visibility into trailer "
                  "availability and last known location.",
        biz_bullets=[],
        impl_items=[
            "Add trailer master records with current location tracking.",
            "When a trailer is selected in dispatch, surface current location immediately.",
            "On dispatch completion, update trailer location to the latest confirmed location.",
            "Provide a trailer list view for operations planning.",
        ],
        ui_items=[
            "Inventory > Dispatch > Trailers",
            "Dispatch Ticket form (trailer and location context)",
        ],
        decisions=[
            "Should trailer location update only at completion, or at intermediate stages as well?",
            "Should one trailer be blocked from assignment to multiple active dispatch tickets?",
        ],
        images=[
            ("01_container_list_proposed_menu_and_columns.png",
             "Proposed menu and dispatch columns on container list"),
            ("02_operations_dropdown_proposed_dispatch_trailers.png",
             "Proposed Dispatch and Trailers menu placement under Operations"),
        ],
    )

    # ── Requirement 4: Manual Dispatch / Route Creation
    req_section(doc, 4, "Manual Dispatch / Route Creation",
        biz_intro="Current behaviour auto-generates routes when parts are created. "
                  "The client requires user-driven creation through explicit actions.",
        biz_bullets=[],
        impl_items=[
            "Disable automatic 3-route generation from parts creation flow.",
            {"text": "Add explicit workflow actions:", "subs": [
                "Create Pickup Dispatch",
                "Create Return Container Dispatch",
                "Create Internal Transfer",
                "Create Delivery Order",
            ]},
            "Enforce duplicate protection for active, equivalent dispatch records.",
            "Preserve full linkage to Sales Order and container record.",
        ],
        ui_items=[
            "Container tracking workflow form header / actions",
            "Dispatch menu for direct manual creation",
        ],
        decisions=[
            "Confirm complete removal of auto 3-route generation.",
            "Confirm whether standalone dispatch creation is allowed from the menu.",
        ],
        images=[
            ("03_container_form_proposed_dispatch_buttons.png",
             "Proposed manual workflow action buttons on container form"),
        ],
    )

    # ── Requirement 5: Internal Transfer Serial Filtering
    req_section(doc, 5, "Internal Transfer Serial Filtering by Sales Order",
        biz_intro="In internal transfers, users need serial selection constrained to serials "
                  "associated with the selected Sales Order.",
        biz_bullets=[],
        impl_items=[
            "Add Sales Order context field to internal transfer form.",
            {"text": "Filter serial / lot selection in Detailed Operations by:", "subs": [
                "Selected Sales Order",
                "Product",
                "Source location availability",
            ]},
            "Display clear blocking / warning behaviour when no valid serials are available.",
            ('Note: The "Not Available" indicator on transfer lines is a stock availability '
             "condition and should be treated separately from filter logic.",
             ['"Not Available"']),
        ],
        ui_items=[
            "Inventory > Operations > Internal Transfers (form)",
            "Detailed Operations modal (serial selection)",
        ],
        decisions=[
            "If Sales Order is not selected, should all serials be visible or selection be "
            "constrained until SO is set?",
            "If SO-linked serials are unavailable, should validation be blocked with a mandatory error?",
        ],
        images=[
            ("04_internal_transfer_list_proposed_so_context.png",
             "Proposed SO context column on internal transfer list"),
            ("05_internal_transfer_form_proposed_so_filter.png",
             "Proposed Sales Order filter field on internal transfer form"),
            ("06_detailed_operations_proposed_so_serial_filter.png",
             "Proposed SO-based serial filtering in Detailed Operations modal"),
        ],
    )

    # ── Delivery Estimate
    page_break(doc)
    add_section_label(doc, "Estimates & Planning")
    heading1(doc, "8. Delivery Estimate (Approved Scope)")
    add_horizontal_rule(doc)
    heading2(doc, "8.1  Effort Breakdown")
    build_effort_table(doc)

    heading2(doc, "8.2  Schedule Estimate")
    estimates = [
        ("Engineering effort", "approximately 91 hours"),
        ("Equivalent duration at 8 hrs/day", "11 to 12 working days"),
        ("Including review, approval, and UAT buffer", "13 to 15 working days"),
    ]
    tbl2 = doc.add_table(rows=len(estimates), cols=2)
    tbl2.style = "Table Grid"
    tbl2.columns[0].width = Inches(3.2)
    tbl2.columns[1].width = Inches(3.0)
    for i, (label, value) in enumerate(estimates):
        bg = LIGHT_BLUE if i % 2 == 0 else WHITE
        lc = tbl2.rows[i].cells[0]
        vc = tbl2.rows[i].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(vc, bg)
        lr = lc.paragraphs[0].add_run(label)
        lr.font.size = Pt(10)
        lr.font.bold = True
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)
        if i == len(estimates) - 1:
            vr.font.bold = True
            vr.font.color.rgb = c(ACCENT_BLUE)

    spacer(doc, 8)

    # ── Phases
    build_phases(doc)

    # ── Risks
    heading1(doc, "10. Delivery Risks and Dependencies")
    add_horizontal_rule(doc)
    body(doc, "Potential factors that may affect the delivery timeline:")
    for risk in [
        "Pending functional decisions listed in Sections 3–7",
        "UAT feedback requiring additional behaviour changes",
        "Data quality issues in stock / serial records in the existing environment",
    ]:
        bullet(doc, risk)

    spacer(doc, 6)

    # ── Out of Scope
    heading1(doc, "11. Out of Scope")
    add_horizontal_rule(doc)
    body(doc, "The following items are explicitly excluded from this estimate:")
    for oos in [
        "WhatsApp API automation / integration",
        "GPS / live telematics integration",
        "Route optimisation engine",
        "Dedicated driver mobile application",
        "Extended multi-company dispatch policy customisation beyond current requirement set",
    ]:
        bullet(doc, oos)

    spacer(doc, 6)

    # ── Executive Summary
    heading1(doc, "12. Executive Summary")
    add_horizontal_rule(doc)

    tbl_ex = doc.add_table(rows=1, cols=1)
    tbl_ex.style = "Table Grid"
    ec = tbl_ex.cell(0, 0)
    set_cell_bg(ec, LIGHT_BLUE)
    set_cell_border(ec,
        left={"val": "single", "sz": "24", "color": h(ACCENT_BLUE)},
        top={"val": "single", "sz": "6", "color": h(ACCENT_BLUE)},
        bottom={"val": "single", "sz": "6", "color": h(ACCENT_BLUE)},
        right={"val": "single", "sz": "6", "color": h(ACCENT_BLUE)},
    )
    ec.paragraphs[0]._element.getparent().remove(ec.paragraphs[0]._element)

    ex_p = ec.add_paragraph()
    ex_p.paragraph_format.left_indent = Inches(0.15)
    ex_p.paragraph_format.space_before = Pt(10)
    ex_p.paragraph_format.space_after = Pt(10)

    r1 = ex_p.add_run(
        "This scope is clearly defined and implementation-ready once the listed decisions are "
        "confirmed. Based on current requirements, delivery is estimated at "
    )
    r1.font.size = Pt(11)
    r1.font.color.rgb = c(DARK_NAVY)

    r2 = ex_p.add_run("13 to 15 working days including UAT buffer")
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.color.rgb = c(ACCENT_BLUE)

    r3 = ex_p.add_run(
        ", using manual WhatsApp dispatch tracking and controlled workflow actions."
    )
    r3.font.size = Pt(11)
    r3.font.color.rgb = c(DARK_NAVY)

    spacer(doc, 8)

    out_path = "/opt/odoo19/e4c/tf_container_tracking/Dispatch_Transfer_Enhancement_Plan.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
